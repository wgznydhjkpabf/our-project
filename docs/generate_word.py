# -*- coding: utf-8 -*-
"""生成校园二手交易平台架构设计说明书 Word 文档"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

OUTPUT = os.path.join(os.path.dirname(__file__), '架构设计说明书.docx')


def set_cell_font(cell, text, bold=False, size=10.5):
    cell.text = ''
    p = cell.paragraphs[0]
    run = p.add_run(str(text))
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        set_cell_font(table.rows[0].cells[i], h, bold=True)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            set_cell_font(table.rows[ri + 1].cells[ci], val)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return table


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = '黑体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        run.font.color.rgb = RGBColor(0, 0, 0)
    return h


def add_para(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(12)
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.bold = bold
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.first_line_indent = Cm(0.74)
    return p


def add_code(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(9)
    run.font.name = 'Consolas'
    p.paragraph_format.line_spacing = 1.2
    p.paragraph_format.left_indent = Cm(0.5)
    return p


def build():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

    # 封面
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run('校园二手交易平台\n架构设计说明书')
    r.font.size = Pt(22)
    r.font.bold = True
    r.font.name = '黑体'
    r._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

    for line in ['', '版本：V2.0', '更新日期：2025-06-14', '技术栈：Java Spring Boot + React + Ant Design + MySQL', '']:
        p = doc.add_paragraph(line)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if p.runs:
            p.runs[0].font.size = Pt(12)

    doc.add_page_break()

    # 一、项目概述
    add_heading(doc, '一、项目概述', 1)
    add_heading(doc, '1.1 项目背景', 2)
    add_para(doc, '本系统面向高校学生群体，提供校园内闲置物品（教材、电子产品、生活用品等）的发布、浏览、交易与管理功能。系统采用 B/S 架构、前后端分离设计，后端提供 RESTful API，前端以 React 单页应用（SPA）呈现图形化操作界面。')

    add_heading(doc, '1.2 设计目标', 2)
    add_table(doc, ['目标', '说明'], [
        ['前后端分离', '前端与后端独立开发、独立部署，通过 HTTP/JSON 通信'],
        ['可维护性', '后端分层架构，前端按页面/组件模块化组织'],
        ['数据规范', '关系数据库满足第三范式（3NF），完整性约束完备'],
        ['安全性', 'JWT 无状态鉴权、密码 BCrypt 加密、权限分级控制'],
        ['可扩展性', '模块化设计，便于后续增加支付、通知推送等功能'],
    ])

    add_heading(doc, '1.3 项目目录结构', 2)
    add_code(doc, '''campus-trade/
├── backend/          # Java 后端（Spring Boot）
│   ├── controller/   # 接口层
│   ├── service/      # 业务层
│   ├── mapper/       # 数据访问层
│   └── entity/       # 实体类
├── frontend/         # React 前端
│   ├── pages/        # 页面
│   ├── components/   # 组件
│   └── api/          # 接口封装
└── sql/              # 数据库脚本''')

    # 二、系统总体架构
    add_heading(doc, '二、系统总体架构', 1)
    add_heading(doc, '2.1 逻辑架构', 2)
    add_para(doc, '系统采用经典三层结构：表现层（React 前端）→ 业务层（Spring Boot 后端）→ 数据层（MySQL 数据库）。前端通过 Axios 发送 RESTful 请求，后端经 JWT 鉴权后由 Controller 调用 Service，Service 通过 MyBatis Mapper 访问数据库。')

    add_code(doc, '''浏览器（React + Ant Design）
        ↓ HTTP/JSON + JWT
Spring Boot（Controller → Service → Mapper）
        ↓ JDBC
MySQL 8.0（7张表 + 存储过程 + 触发器）''')

    add_heading(doc, '2.2 部署架构', 2)
    add_table(doc, ['组件', '端口', '说明'], [
        ['React 前端', '5173', 'Vite 开发服务器 / 生产环境 Nginx'],
        ['Spring Boot 后端', '8080', 'REST API + 文件上传'],
        ['MySQL', '3306', '数据持久化'],
    ])

    add_heading(doc, '2.3 技术架构', 2)
    add_table(doc, ['层次', '技术'], [
        ['前端', 'React + Ant Design'],
        ['后端', 'Java（Spring Boot）'],
        ['数据库', 'MySQL'],
        ['架构', 'B/S，前后端分离'],
    ])

    add_heading(doc, '2.4 技术选型详表', 2)
    add_table(doc, ['层次', '技术', '版本', '说明'], [
        ['后端框架', 'Spring Boot', '3.2.5', 'REST 服务'],
        ['安全框架', 'Spring Security + JWT', '—', '无状态鉴权'],
        ['持久层', 'MyBatis', '3.0.3', 'SQL 映射'],
        ['前端框架', 'React', '18.3', 'SPA 组件化'],
        ['UI 库', 'Ant Design', '5.x', '企业级组件'],
        ['路由', 'React Router', '6.x', '前端路由'],
        ['HTTP', 'Axios', '1.x', 'API 请求'],
        ['构建', 'Vite', '5.x', '前端构建'],
        ['数据库', 'MySQL', '8.0', '关系型数据库'],
    ])

    # 三、功能模块
    add_heading(doc, '三、功能模块设计', 1)
    add_heading(doc, '3.1 功能模块', 2)
    add_code(doc, '''校园二手交易平台
├── 用户模块      注册 / 登录 / 个人资料
├── 商品模块      发布 / 浏览 / 搜索 / 下架
├── 订单模块      下单 / 确认 / 完成 / 取消
├── 评价模块      评分 / 信誉分
├── 消息模块      站内私信
├── 地址模块      收货地址管理
└── 管理模块      审核 / 用户管理 / 分类 / 统计''')

    add_heading(doc, '3.2 后端模块映射', 2)
    add_table(doc, ['模块', 'Controller', 'Service'], [
        ['认证/用户', 'AuthController', 'AuthService'],
        ['商品', 'GoodsController', 'GoodsService'],
        ['订单/评价', 'OrderController', 'OrderService'],
        ['消息/地址/分类', 'CommonController', 'CommonService'],
        ['管理后台', 'AdminController', 'GoodsService / CommonService'],
        ['文件上传', 'UploadController', '—'],
    ])

    add_heading(doc, '3.3 前端页面映射', 2)
    add_table(doc, ['模块', '页面', '路由'], [
        ['首页', 'Home.jsx', '/'],
        ['登录/注册', 'Login.jsx / Register.jsx', '/login / /register'],
        ['商品详情', 'GoodsDetail.jsx', '/goods/:id'],
        ['发布/我的商品', 'Publish.jsx / MyGoods.jsx', '/publish / /my-goods'],
        ['订单', 'Orders.jsx', '/my-orders'],
        ['个人中心', 'Profile.jsx', '/profile'],
        ['消息', 'Messages.jsx', '/messages'],
        ['管理后台', 'Admin.jsx', '/admin'],
    ])

    # 四、后端架构
    add_heading(doc, '四、后端架构设计', 1)
    add_heading(doc, '4.1 分层架构', 2)
    add_code(doc, '请求 → JwtAuthFilter → Controller → Service → Mapper → MySQL')
    add_table(doc, ['层次', '职责', '示例'], [
        ['Controller', '接收请求、参数校验、返回 JSON', 'AuthController'],
        ['Service', '业务逻辑、事务控制', 'OrderService'],
        ['Mapper', '数据访问（MyBatis）', 'UserMapper'],
        ['Entity', '数据库实体', 'User, Goods'],
        ['DTO', '请求传输对象', 'RegisterRequest'],
        ['Config', '安全、JWT、跨域', 'SecurityConfig'],
    ])

    add_heading(doc, '4.2 安全设计', 2)
    add_para(doc, '用户登录成功后，后端使用 BCrypt 校验密码并签发 JWT（含 userId、role，有效期 24 小时）。前端将 Token 存入 localStorage，后续请求在 Header 中携带 Authorization: Bearer <token>。JwtAuthFilter 解析 Token 并注入 SecurityContext。')
    add_table(doc, ['接口类型', '策略'], [
        ['公开接口', '/api/auth/**、商品浏览、分类列表'],
        ['登录用户', '发布商品、下单、消息、地址等'],
        ['管理员', '/api/admin/**，校验 role == 1'],
    ])

    add_heading(doc, '4.3 订单状态机', 2)
    add_table(doc, ['状态值', '含义', '说明'], [
        ['0', '待确认', '买家下单后，等待卖家确认'],
        ['1', '待交易', '卖家确认，等待线下/邮寄交易'],
        ['2', '已完成', '交易完成，买家可评价'],
        ['3', '已取消', '买卖双方取消订单'],
    ])

    add_heading(doc, '4.4 商品状态', 2)
    add_table(doc, ['状态值', '含义', '触发条件'], [
        ['0', '待审核', '用户发布商品'],
        ['1', '在售', '管理员审核通过'],
        ['2', '已售', '订单完成'],
        ['3', '已下架', '卖家下架或管理员驳回'],
    ])

    # 五、前端架构
    add_heading(doc, '五、前端架构设计', 1)
    add_para(doc, '前端采用 React 18 单页应用，使用 React Router 6 管理路由，Ant Design 5 构建 UI。登录态通过 localStorage 存储 Token，PrivateRoute 组件实现路由守卫。')
    add_table(doc, ['路由', '权限', '说明'], [
        ['/', '公开', '商品列表'],
        ['/login, /register', '公开', '登录注册'],
        ['/goods/:id', '公开', '商品详情'],
        ['/publish 等', '需登录', 'PrivateRoute 守卫'],
        ['/admin', '管理员', 'PrivateRoute admin'],
    ])
    add_para(doc, 'HTTP 请求统一封装在 api/request.js（Axios 实例），请求拦截器注入 JWT，响应拦截器解析统一 Result 格式并弹出错误提示。')

    # 六、数据库设计
    add_heading(doc, '六、数据库架构设计', 1)
    add_heading(doc, '6.1 数据表（7张，满足3NF）', 2)
    add_table(doc, ['表名', '说明', '主要外键'], [
        ['sys_user', '用户表', '—'],
        ['goods_category', '商品分类', '—'],
        ['goods', '商品表', 'user_id, category_id'],
        ['trade_order', '订单表', 'goods_id, buyer_id, seller_id'],
        ['order_review', '评价表', 'order_id, reviewer_id'],
        ['user_message', '消息表', 'sender_id, receiver_id'],
        ['user_address', '地址表', 'user_id'],
    ])

    add_heading(doc, '6.2 完整性约束', 2)
    add_table(doc, ['类型', '示例'], [
        ['主键', '每表 AUTO_INCREMENT 主键'],
        ['唯一', '学号、邮箱、订单号唯一'],
        ['外键', 'goods→user, order→goods'],
        ['CHECK', 'price > 0, rating 1-5'],
    ])

    add_heading(doc, '6.3 存储过程与触发器', 2)
    add_para(doc, '存储过程 sp_complete_order：订单完成时更新订单状态、商品为已售、重算卖家信誉分。')
    add_para(doc, '触发器 trg_order_status_before_update：禁止已完成订单被修改状态。')

    # 七、接口设计
    add_heading(doc, '七、接口设计（RESTful API）', 1)
    add_para(doc, '基础路径：http://localhost:8080/api，数据格式 application/json，认证方式 Authorization: Bearer <JWT>。')

    add_heading(doc, '7.1 认证模块 /api/auth', 2)
    add_table(doc, ['方法', '路径', '权限', '说明'], [
        ['POST', '/register', '公开', '用户注册'],
        ['POST', '/login', '公开', '登录返回 JWT'],
        ['GET', '/profile', '登录', '获取个人信息'],
        ['PUT', '/profile', '登录', '更新个人信息'],
    ])

    add_heading(doc, '7.2 商品模块 /api/goods', 2)
    add_table(doc, ['方法', '路径', '权限', '说明'], [
        ['GET', '/', '公开', '商品列表'],
        ['GET', '/{id}', '公开', '商品详情'],
        ['POST', '/', '登录', '发布商品'],
        ['PUT', '/{id}', '登录', '编辑商品'],
        ['PUT', '/{id}/off-shelf', '登录', '下架商品'],
    ])

    add_heading(doc, '7.3 订单模块 /api/orders', 2)
    add_table(doc, ['方法', '路径', '权限', '说明'], [
        ['POST', '/', '登录', '创建订单'],
        ['GET', '/', '登录', '我的订单'],
        ['PUT', '/{id}/confirm', '登录', '卖家确认'],
        ['PUT', '/{id}/cancel', '登录', '取消订单'],
        ['PUT', '/{id}/complete', '登录', '完成订单'],
        ['POST', '/review', '登录', '提交评价'],
    ])

    add_heading(doc, '7.4 管理模块 /api/admin', 2)
    add_table(doc, ['方法', '路径', '说明'], [
        ['GET', '/dashboard', '统计数据'],
        ['GET', '/users', '用户列表'],
        ['GET', '/goods/pending', '待审核商品'],
        ['PUT', '/goods/{id}/audit', '审核商品'],
        ['POST/PUT/DELETE', '/categories', '分类管理'],
    ])

    # 八、业务流程
    add_heading(doc, '八、关键业务流程', 1)
    add_para(doc, '完整交易流程：买家浏览商品 → 查看详情 → 创建订单 → 卖家确认 → 线下/邮寄交易 → 双方确认完成（调用 sp_complete_order）→ 买家评价 → 更新卖家信誉分。')
    add_para(doc, '商品审核流程：卖家发布（待审核）→ 管理员审核 → 通过（在售）/ 驳回（下架）。')

    # 九、部署环境
    add_heading(doc, '九、开发与部署环境', 1)
    add_table(doc, ['工具', '版本'], [
        ['JDK', '17+'],
        ['Maven', '3.8+'],
        ['Node.js', '18+'],
        ['MySQL', '8.0+'],
    ])
    add_para(doc, '数据库配置：url=jdbc:mysql://localhost:3306/campus_trade，username=root，password=123456。')
    add_para(doc, '启动顺序：① MySQL 启动并执行 sql/init.sql → ② 后端 mvn spring-boot:run（8080）→ ③ 前端 npm run dev（5173）。')
    add_para(doc, '默认管理员账号：admin / admin123（首次启动自动创建）。')

    doc.save(OUTPUT)
    print(f'已生成: {OUTPUT}')


if __name__ == '__main__':
    build()
